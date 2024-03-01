import docx
import os
import PyPDF2
import re
import spacy
import subprocess
import sqlite3


# Download spaCy model if not already installed
try:
    nlp = spacy.load("en_core_web_sm")
except OSError:
    subprocess.run(["python", "-m", "spacy", "download", "en_core_web_sm"])
    nlp = spacy.load("en_core_web_sm")


def get_data_directory_path():
    """
        function is used to get path
    """
    ROOT_DIR = os.getcwd()
    # This is to extract the patent directory from the ful path(ETLSystem)
    ROOT_DIR = os.path.dirname(ROOT_DIR)
    DATA_DIR = os.path.join(ROOT_DIR, "rawdata")

    return DATA_DIR


def word_extractor(filepath):
    """
        function is used to extract data from a word document
    """
    doc = docx.Document(filepath)  # Reading in the file using Document method
    extracted_text = []  # an array that would hold the extracted file
    for paragraph in doc.paragraphs:
        # loop through each paragraph in the word doc and extracting the text.
        extracted_text.append(paragraph.text)

    return "\n".join(extracted_text)


def pdf_extractor(filepath):
    """
        function is used to extract data from a word pdf document
    """
    with open(filepath, 'rb') as pdf_file:  # reading in the pdf file
        # using a pdfreader to read in the file as a pdf
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        extracted_text = ""  # empty string that holds the extracted data

        # loops that loop over each page of the pdf and extract text form it and add its to the empty string
        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            page_text = page.extract_text()
            extracted_text += page_text

    return extracted_text


def mask_personal_information_2(text):
    """
    Takes text and redacts "PERSON", "GPE", "DATE", "PHONE", "NORP", "ORG","EMAIL", "LOC", "FAC" from it
    :param str text: text to be redacted
    :return: redacted text
    """
    doc = nlp(text)
    redacted_text = text

    # Redaction with spacy
    for ent in doc.ents:
        if ent.label_ in ["PERSON", "GPE", "DATE", "PHONE", "NORP", "ORG", "EMAIL", "LOC", "FAC"]:
            # Redact entities like names, organizations, locations, and dates
            redacted_text = redacted_text.replace(ent.text, "****")

    # Add more redaction for properties not caught by spacy
    # Redact email addresses using regex
    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,4}\b'
    redacted_text = re.sub(email_pattern, "*****", redacted_text)

    # Define a regular expression pattern to match both common and additional phone number formats
    phone_number_pattern = r'\+?\d{0,4}\s?\(?\d+\)?\s?\d+\s?\d+\s?\d+|\(\d{3}\)\s?\d{3}-\d{4}'
    redacted_text = re.sub(phone_number_pattern,
                           "*******", redacted_text)

    # Define a regular expression pattern to match both common and additional phone number formats
    name_pattern = r'\b[A-Z][A-Za-z]* [A-Z][A-Za-z]* [A-Z][A-Za-z]*\b'
    redacted_text = re.sub(
        name_pattern, "*******", redacted_text)

    return redacted_text


def find_emails(text):
    # Regular expression pattern for matching email addresses
    pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'

    # Using re.findall() to find all email addresses in the text
    emails = re.findall(pattern, text)

    return emails


def fetch_maskedcv_jobdes(email):
    conn = sqlite3.connect('mydatabase.db')
    cursor = conn.cursor()

    cursor.execute(
        "SELECT maskedcv, maskedjobdes FROM CareerCraft WHERE useremail=?", (email,))

    result = cursor.fetchone()

    conn.close()

    if result:
        return result[0], result[1]
    else:
        return None, None


def save_text_as_word(text, filename):
    # Create a new Document
    doc = docx.Document()

    # Add the text to the document
    doc.add_paragraph(text)

    # Save the document
    doc.save(filename)
